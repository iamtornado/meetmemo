"""Build Word (.docx) export for meeting transcripts with speaker and timestamps."""

from __future__ import annotations

import re
from datetime import date, datetime
from io import BytesIO

from docx import Document
from app.models.transcript import Transcript, TranscriptSegment


def _format_time(seconds: float) -> str:
    total = max(0, int(seconds))
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def _speaker_label(seg: TranscriptSegment, *, unknown: str = "—") -> str:
    if seg.speaker_name:
        return seg.speaker_name
    if seg.speaker_id:
        return seg.speaker_id
    return unknown


def build_transcript_docx(
    *,
    title: str,
    transcript: Transcript,
    meeting_date: date | datetime | None = None,
) -> bytes:
    doc = Document()
    doc.add_heading(title or "Meeting Transcript", level=0)

    meta_parts: list[str] = []
    if meeting_date:
        if isinstance(meeting_date, datetime):
            meeting_date = meeting_date.date()
        meta_parts.append(f"Date: {meeting_date.isoformat()}")
    if transcript.language:
        meta_parts.append(f"Language: {transcript.language}")
    if transcript.model_used:
        meta_parts.append(f"ASR model: {transcript.model_used}")
    meta_parts.append(f"Segments: {len(transcript.segments)}")
    if transcript.word_count:
        meta_parts.append(f"Word count: {transcript.word_count}")

    if meta_parts:
        doc.add_paragraph(" · ".join(meta_parts))

    segments = sorted(transcript.segments, key=lambda s: s.seq_number)
    if not segments:
        doc.add_paragraph("No transcript segments.")
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    lang = (transcript.language or "").lower()
    zh = lang.startswith("zh")
    if zh:
        headers = ("开始时间", "结束时间", "说话人", "内容")
        unknown_speaker = "未知"
    else:
        headers = ("Start", "End", "Speaker", "Text")
        unknown_speaker = "—"
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for seg in segments:
        row = table.add_row().cells
        row[0].text = _format_time(seg.start_time)
        row[1].text = _format_time(seg.end_time)
        row[2].text = _speaker_label(seg, unknown=unknown_speaker)
        row[3].text = (seg.text or "").strip()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def transcript_export_filename(title: str | None, meeting_id: str) -> str:
    base = (title or "transcript").strip()
    base = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", base)
    base = re.sub(r"\s+", " ", base).strip("._ ") or "transcript"
    if len(base) > 80:
        base = base[:80].rstrip("._ ")
    return f"{base}.docx"
