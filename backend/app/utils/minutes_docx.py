"""Export formal meeting minutes text to Word (.docx)."""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def build_formal_minutes_docx(title: str, formal_minutes: str) -> bytes:
    doc = Document()
    lines = [ln.rstrip() for ln in formal_minutes.splitlines()]

    title_set = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if not title_set and stripped in ("纪要", "会 纪 要"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("纪 要")
            run.bold = True
            run.font.size = Pt(22)
            title_set = True
            continue

        if stripped.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、")):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            continue

        if stripped.startswith(("会议主题", "会议时间", "会议地点", "会议主持", "会议记录", "纪要审核", "出席")):
            doc.add_paragraph(stripped)
            continue

        doc.add_paragraph(stripped)

    if not title_set:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(title or "会议纪要")
        run.bold = True
        run.font.size = Pt(18)
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def minutes_export_filename(title: str | None) -> str:
    base = (title or "会议纪要").strip()
    base = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", base)
    base = re.sub(r"\s+", " ", base).strip("._ ") or "会议纪要"
    if len(base) > 80:
        base = base[:80].rstrip("._ ")
    return f"{base}-纪要.docx"
